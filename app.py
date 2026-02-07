
import streamlit as st
from openai import OpenAI
from PyPDF2 import PdfReader
from docx import Document
import os
import re
from dotenv import load_dotenv
load_dotenv()
api_key = st.secret["OPENAI_API_KEY"] if "OPENAI_APIA-KEY" in st.secrets else os.getenv("OPENAI_API_KEY")
client = OpenAI(api_key=api_key)
if "resume_text" not in st.session_state:
    st.session_state.setdefault("resume_text", "")
st.title("🚀AI Resume Analyzer - Get Recruiter-Level Feedback in seconds")
st.caption("Built by an AI Engineer - Available for freelance AI projects.")
st.divider()
uploaded_file = st.file_uploader("📑Upload your resume", type=["pdf", "docx"])
job_description = st.text_area(
    "💼 Paste the job description (optional, but highly recommended)",
    height=200,
    placeholder="Paste the job posting here to get a highly targeted resume rewrite..."
)
resume_text = ""

if uploaded_file is not None:
    #---- PDF ----
    if uploaded_file.type == "application/pdf":
        reader = PdfReader(uploaded_file)
        for page in reader.pages:
            text = page.extract_text()
            if text:
                resume_text += text
    #---- DOCX ----
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = Document(uploaded_file)
        for para in doc.paragraphs:
            if para.text:
                resume_text += para.text + "\n"
                #SAVE IT (critical)
                st.session_state["resume_text"] = resume_text
        
# BUTTON LIVES HERE (not nested)
if st.button("Analyze Resume") or "resume_text" in st.session_state:
    with st.spinner("🔍 AI is analyzing your resume... This takes about 5-10 seconds"):
        response = client.chat.completions.create(
            model="gpt-4.1-mini",
             messages=[
{
        "role": "system",
        "content": """You are an elite technical recruiter and resume coach.

    Analyze the resume and return feedback using this EXACT format with emojis:
    if a job description is provided:
    ⭐️ Give a match score percentage
    🎯Identify missing keywords
    🚀 Recommend ways to increase the match score
    📈 Suggest targeted improvements

    CRITICAL INSTRUCTIONS:
    The fIRST LINE of your response Must be formatted like this:
OverallScore: 78
- No bold
- No bullets
- No extra words
- No emojis on this line
Then continue with the resume analysis
    
    
  ✅ Strengths
- List strong points

    ⚠️ Weaknesses
    - List improvement areas

    🚀 How To Improve Immediately
    - Give actionable upgrades

    🎯 ATS Optimization Tips
    - Suggest keywords and formatting improvements

    Keep feedback clear, modern, and highly actionable.
    Avoid long paragraphs.
    Use bullet points.
    Sound like a recruiter that candidates would pay for.
"""
    },
                {
                    "role": "user",
                    "content": f"""
                    Rewrite this resume specifically for the job below.
                    JOB DESCRIPTION: 
                    {job_description}  
                    RESUME:
                    {st.session_state.get("resume_text", "")}
                    """
                }
            ],
            max_tokens=700,
            temperature=0.3,
        )
    result = response.choices[0].message.content
    st.subheader("📝 Recruiter-Level Resume Breakdown")
    st.markdown(response.choices[0].message.content)
    st.session_state["resume_text"] = resume_text
    
    # Extract score from AI response
    match = re.search(r"overall.*?(\d{1,3})", result, re.IGNORECASE)
    if match:
        score = int(match.group(1))
        st.session_state["score"] = score
        st.divider()
        st.subheader("📊 Resume Strength Score")
        st.progress(score / 100)
        st.metric("Overall Score", f"{score}/100")
    else:
        st.info("Score not detected. Try again")
    
    # HIGH-VALUE FEATURE - REWRITE RESUME BUTTON
    st.divider()
    
if st.session_state.get("resume_text") and st.session_state.get("score"):
        #Prevent empty rewrites
        if not st.session_state.get("resume_text", ""):
            st.warning("⚠️ Please upload a resume first.")
            st.stop()
        if st.button("🚀 Rewrite Resume for This Job"):
            with st.spinner("✍️ AI is rewriting your resume... This takes about 5-10 seconds"):
                old_score = st.session_state.get("score", 0)
                st.write("Resume length", len(st.session_state.get("resume_text", "")))
                rewrite_response = client.chat.completions.create(
                    model="gpt-4.1-mini",
                    messages=[
                        {
                            "role": "system",
                            "content": """
                            Your primary objective is to increase the ATS match score.
                            Do not priortize sounding fancy
                            Prioritize keyword aalignment with the job description
                            Maintain or increase keyword density
                            Do not remove relevant technical skills
                            Mirror the language used in job posting.
                            if the resume already contains strong content, improve it without removing important keywords.
                                You are an elite resume writer and career coach. 

                                Rewrite the resume to be: 
                                CRITICAL :Optimize  the resume specifically to maximize the ATS match score for the provided job description.Mirror important keywords from the job description naturally.
                                Result-driven
                                ATS-optimized
                                Modern
                                Professional
                                Achievment-focused
                                Keyword-rich (but natural)
                                Improve bullet points
                                Add strong measurable achievements(%,$,time saved,revenue,performance improvements).
                                Make it sound like a top 10% candidate in this field
                                Return the FULL rewritten resume
                                Optimize the resume specifically for the provided job description.
                                Increase the match score as much as possible
                                Mirror important keywords from the job posting naturally
                                Return a HIGHLY optimized resume designed to score above 85%.
                            """
                        },
                        {
                            "role": "user",
                            "content":f"""
                            JOB DESCRIPTION:
                            {job_description} 
                            RESUME:
                            {st.session_state.get("resume_text", "")}
                             """
                             }
                    ],
                    max_tokens=1200,
                    temperature=0.3,
                )
                rewritten_resume = rewrite_response.choices[0].message.content
                st.session_state["resume_text"] = rewritten_resume
                #RE-SCORE THE NEW RESUME
                analysis_response = client.chat.completions.create(
                    model="gpt-4.1-mini",
                    messages=[
                        {
                            "role": "system",
                            "content":"""
                            You are an elite technical recruiter
                            score this resume from 1-100
                            IMPORTANT:
                            A professionally rewritten  resume should NEVER score below the original resume. If the rewritten resume is not better, return a score that is the same or higher than the original resume.
                            FIRST LINE MUST BE
                            OverallScore: number
                            No emojis
                            No extra words
                            """
                        },
                        {
                            "role": "user",
                            "content":rewritten_resume
                        }
                    ],
                    max_tokens=500,
                    temperature=0.2
                    )
                
                analysis_result = analysis_response.choices[0].message.content
                import re
                match = re.search(r'(\d{1,3})', analysis_result)
                if match:
                    new_score = int(match.group(1))
                    st.session_state["score"] = new_score
                    score= new_score
                    st.divider()
                    st.subheader("📈 New Resume Improvment")
                    col1, col2, col3 = st.columns(3)
                    col1.metric("Old Score", f"{old_score}/100")
                    col2.metric("New Score", f"{new_score}/100")
                    increase = new_score - old_score
                    col3.metric("Increase", f"{increase}",delta=increase)
                resume_text = rewritten_resume
                rewritten_resume=rewritten_resume.encode("latin-1", "replace").decode("latin-1")
                st.write(rewritten_resume)
                st.divider()
                #create Word document
                document = Document()
                document.add_heading('Rewritten Resume ', level=1)
                clean_text = rewritten_resume.replace("**", "").replace("*", "")
                for line in clean_text.split('\n'):
                    document.add_paragraph(line)
                    #Save to memory
                    from io import BytesIO
                docx_buffer = BytesIO()
                document.save(docx_buffer)
                docx_buffer.seek(0)
                st.download_button(
                    label="📥 Download Rewritten Resume",
                    data=docx_buffer,
                    file_name="rewritten_resume.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                st.divider()
                st.success("🔥 Your Resume has been upgraded")
        st.info("💡 Tip: Tailor your resume for every job. ATS systems filter resumes before a human ever sees them.")